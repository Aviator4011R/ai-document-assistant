import os
import tempfile
from flask import Blueprint, request, jsonify
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import json

document_bp = Blueprint('document', __name__)

# Global variables to store the current document's vector store and metadata
current_vector_store = None
current_document_title = None
current_document_content = None

# Initialize OpenAI components with proper error handling
embeddings = None
llm = None

def initialize_openai_services():
    """Initialize OpenAI services with error handling"""
    global embeddings, llm
    
    try:
        from langchain_openai import OpenAIEmbeddings, ChatOpenAI
        from langchain_community.vectorstores import FAISS
        from langchain.chains import RetrievalQA
        
        # Check if API key is available
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key or len(api_key) < 20:
            print("Warning: OpenAI API key not properly configured")
            return False
        
        # Use supported models
        embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
        llm = ChatOpenAI(model="gpt-4.1-mini", temperature=0)
        
        print("OpenAI services initialized successfully")
        return True
        
    except Exception as e:
        print(f"Warning: OpenAI initialization failed: {e}")
        # Try alternative approach with simple text processing
        try:
            global simple_mode
            simple_mode = True
            print("Falling back to simple text processing mode")
            return True
        except:
            return False

# Try to initialize on import
simple_mode = False
initialize_openai_services()

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")
    return text

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except Exception as e:
        try:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
        except Exception as e2:
            raise Exception(f"Error reading TXT: {str(e2)}")
    return text

def process_document_text(text, title):
    """Process document text into vector store or simple storage"""
    global current_vector_store, current_document_title, current_document_content
    
    if simple_mode or not embeddings or not llm:
        # Simple mode - just store the text
        current_document_content = text
        current_document_title = title
        current_vector_store = "simple_mode"  # Flag to indicate document is loaded
        
        # Split into chunks for better searching
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        chunks = text_splitter.split_text(text)
        return len(chunks)
    
    else:
        # Full vector mode
        from langchain_community.vectorstores import FAISS
        
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        # Create documents
        documents = [LangchainDocument(page_content=text, metadata={"source": title})]
        texts = text_splitter.split_documents(documents)
        
        # Create vector store
        vector_store = FAISS.from_documents(texts, embeddings)
        
        # Store globally
        current_vector_store = vector_store
        current_document_title = title
        current_document_content = text
        
        return len(texts)

def simple_search(query, text, max_results=3):
    """Simple text search when vector search is not available"""
    query_lower = query.lower()
    text_lower = text.lower()
    
    # Split text into sentences
    sentences = text.split('.')
    
    # Find sentences containing query terms
    relevant_sentences = []
    query_words = query_lower.split()
    
    for sentence in sentences:
        sentence_lower = sentence.lower()
        score = sum(1 for word in query_words if word in sentence_lower)
        if score > 0:
            relevant_sentences.append((sentence.strip(), score))
    
    # Sort by relevance and return top results
    relevant_sentences.sort(key=lambda x: x[1], reverse=True)
    return [s[0] for s in relevant_sentences[:max_results]]

@document_bp.route('/upload', methods=['POST'])
def upload_document():
    """Upload and process a document"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed. Please upload PDF, DOCX, or TXT files.'}), 400
        
        # Save file temporarily
        filename = secure_filename(file.filename)
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, filename)
        file.save(file_path)
        
        try:
            # Extract text based on file type
            file_extension = filename.rsplit('.', 1)[1].lower()
            
            if file_extension == 'pdf':
                text = extract_text_from_pdf(file_path)
            elif file_extension == 'docx':
                text = extract_text_from_docx(file_path)
            elif file_extension == 'txt':
                text = extract_text_from_txt(file_path)
            else:
                return jsonify({'error': 'Unsupported file type'}), 400
            
            # Clean up temporary file
            os.remove(file_path)
            os.rmdir(temp_dir)
            
            if not text.strip():
                return jsonify({'error': 'No text could be extracted from the document'}), 400
            
            # Process the document
            title = filename.rsplit('.', 1)[0]  # Remove extension for title
            chunk_count = process_document_text(text, title)
            
            return jsonify({
                'success': True,
                'title': title,
                'chunk_count': chunk_count,
                'text_length': len(text),
                'message': f'Document "{title}" processed successfully',
                'mode': 'simple' if simple_mode else 'vector'
            })
            
        except Exception as e:
            # Clean up on error
            if os.path.exists(file_path):
                os.remove(file_path)
            if os.path.exists(temp_dir):
                os.rmdir(temp_dir)
            raise e
            
    except Exception as e:
        return jsonify({'error': f'Error processing document: {str(e)}'}), 500

@document_bp.route('/query', methods=['POST'])
def query_document():
    """Query the processed document"""
    global current_vector_store, current_document_title, current_document_content
    
    try:
        if not current_vector_store:
            return jsonify({'error': 'No document has been uploaded and processed'}), 400
        
        data = request.get_json()
        if not data or 'question' not in data:
            return jsonify({'error': 'No question provided'}), 400
        
        question = data['question'].strip()
        if not question:
            return jsonify({'error': 'Question cannot be empty'}), 400
        
        if simple_mode or current_vector_store == "simple_mode":
            # Simple text search mode
            relevant_text = simple_search(question, current_document_content)
            
            if relevant_text:
                answer = f"Based on the document '{current_document_title}', here are the relevant sections:\n\n"
                answer += "\n\n".join(relevant_text)
            else:
                answer = "That information is not part of this document. Please ask a question based on the uploaded manual."
            
            return jsonify({
                'success': True,
                'answer': answer,
                'document_title': current_document_title,
                'mode': 'simple_search'
            })
        
        else:
            # Vector search mode
            if not llm:
                return jsonify({'error': 'AI services not available'}), 500
            
            from langchain.chains import RetrievalQA
            
            # Create retrieval QA chain
            qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=current_vector_store.as_retriever(search_kwargs={"k": 3}),
                return_source_documents=True
            )
            
            # Custom prompt to ensure document-only responses
            custom_prompt = f"""
            You are an AI assistant that answers questions based ONLY on the provided document: "{current_document_title}".
            
            IMPORTANT RULES:
            1. Only answer questions using information that is explicitly contained in the provided document
            2. If the information is not in the document, respond with: "That information is not part of this document. Please ask a question based on the uploaded manual."
            3. Do not use any external knowledge or make assumptions
            4. Be specific and cite relevant sections when possible
            5. Keep responses clear and concise
            
            Question: {question}
            """
            
            # Get response
            result = qa_chain({"query": custom_prompt})
            answer = result['result']
            
            # Check if the answer indicates information not found
            not_found_phrases = [
                "not mentioned", "not provided", "not specified", "not found",
                "does not contain", "no information", "not available"
            ]
            
            if any(phrase in answer.lower() for phrase in not_found_phrases):
                answer = "That information is not part of this document. Please ask a question based on the uploaded manual."
            
            return jsonify({
                'success': True,
                'answer': answer,
                'document_title': current_document_title,
                'sources_used': len(result['source_documents']),
                'mode': 'vector_search'
            })
        
    except Exception as e:
        return jsonify({'error': f'Error processing query: {str(e)}'}), 500

@document_bp.route('/status', methods=['GET'])
def get_status():
    """Get current document processing status"""
    global current_document_title, current_vector_store
    
    return jsonify({
        'has_document': current_vector_store is not None,
        'document_title': current_document_title,
        'ready_for_queries': current_vector_store is not None,
        'ai_services_available': embeddings is not None and llm is not None,
        'mode': 'simple' if simple_mode else 'vector'
    })

@document_bp.route('/clear', methods=['POST'])
def clear_document():
    """Clear the current document and reset the system"""
    global current_vector_store, current_document_title, current_document_content
    
    current_vector_store = None
    current_document_title = None
    current_document_content = None
    
    return jsonify({
        'success': True,
        'message': 'Document cleared successfully'
    })

@document_bp.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'services': {
            'embeddings': embeddings is not None,
            'llm': llm is not None,
            'document_loaded': current_vector_store is not None,
            'simple_mode': simple_mode
        }
    })

